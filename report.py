"""
report.py — Executive .docx report generator for TelecomLens.
Uses python-docx only — no Node.js required.
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Brand colours ────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x25, 0x63, 0xEB)   # #2563EB
TEAL   = RGBColor(0x0D, 0x94, 0x88)   # #0D9488
GRAY   = RGBColor(0x49, 0x50, 0x57)   # #495057
LGRAY  = RGBColor(0xDE, 0xE2, 0xE6)   # #DEE2E6 header fill
BLACK  = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
RED    = RGBColor(0xDC, 0x26, 0x26)   # #DC2626 anomaly
AMBER  = RGBColor(0xD9, 0x77, 0x06)   # #D97706


def _hex(rgb: RGBColor) -> str:
    return str(rgb)  # RGBColor.__str__ returns e.g. '2563EB'


def _set_cell_bg(cell, hex_colour: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def _add_para(doc, text="", style=None, bold=False, size=None, colour=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        if size:   run.font.size = Pt(size)
        if colour: run.font.color.rgb = colour
    return p


def _header_row(table, labels: list[str], widths_cm: list[float]):
    row = table.rows[0]
    for i, lbl in enumerate(labels):
        cell = row.cells[i]
        cell.width = Cm(widths_cm[i])
        _set_cell_bg(cell, _hex(BLUE))
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(lbl)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _data_row(table, values: list, widths_cm: list, shade=False, bold_first=False, red_last=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.width = Cm(widths_cm[i])
        if shade:
            _set_cell_bg(cell, "F8F9FA")
        align = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        para = cell.paragraphs[0]
        para.alignment = align
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(str(val) if val is not None else "—")
        run.font.size = Pt(9)
        if i == 0 and bold_first:
            run.bold = True
        if i == len(values) - 1 and red_last and val and str(val).startswith("▲"):
            run.font.color.rgb = RED
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _section_heading(doc, number: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), _hex(BLUE))
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)
    run_num = p.add_run(number + " ")
    run_num.bold = True
    run_num.font.size = Pt(13)
    run_num.font.color.rgb = BLUE
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(13)
    run_title.font.color.rgb = BLACK
    return p


def _kes(n) -> str:
    try:
        return f"KES {float(n):,.0f}"
    except (TypeError, ValueError):
        return "—"


def _clip(text, n: int = 34) -> str:
    t = str(text or "")
    return t if len(t) <= n else t[:n - 1] + "…"


def generate_report(data: dict) -> bytes:
    """
    Generate a professional executive .docx report from bill data.

    data keys:
      org_name, account_number, statement_date,
      summary (dict with account_total, pre_tax_total, excise_total,
               vat_total, outstanding_total, subscriber_count, anomaly_count),
      divisions (list of {division, total, count}),
      anomalies (list of {subscriber_number, raw_name, division,
                          amount_due_kes, anomaly_reason, cdr_count}),
      top_subscribers (list of {subscriber_number, raw_name, division,
                                 tariff_plan, amount_due_kes}),
      trends (list of {statement_date, account_total, subscriber_count,
                        anomaly_count}),
    """
    s = data.get("summary", {})
    org_name = data.get("org_name") or s.get("org_name", "Unknown Organisation")
    acct = data.get("account_number") or s.get("account_number", "")
    stmt_date = data.get("statement_date") or s.get("statement_date", "")
    generated = datetime.utcnow().strftime("%d %B %Y")

    divisions = data.get("divisions", [])
    anomalies = data.get("anomalies", [])
    top_subs = data.get("top_subscribers", [])[:15]
    trends = data.get("trends", [])

    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default font ────────────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    # ════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════════════
    _add_para(doc, space_before=40)
    _add_para(doc, "TELECOMLENS", bold=True, size=28, colour=BLUE,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_para(doc, "Executive Bill Analysis Report", size=13, colour=GRAY,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    _add_para(doc, org_name, bold=True, size=18, colour=BLACK,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    if acct:
        _add_para(doc, f"Account: {acct}", size=10, colour=GRAY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    if stmt_date:
        _add_para(doc, f"Statement Period: {stmt_date}", size=10, colour=GRAY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_para(doc, f"Report Generated: {generated}", size=9, colour=GRAY,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=0)
    _add_para(doc, "Confidential — For Internal Use Only", size=9, colour=GRAY,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "1.", "Executive Summary")
    _add_para(doc,
        f"This report analyses the Safaricom postpay bill for {org_name} "
        f"(account {acct or '—'}) for the billing period ending {stmt_date}. "
        f"The total amount due is {_kes(s.get('account_total'))} covering "
        f"{s.get('subscriber_count', 0)} active subscriber lines. "
        f"{s.get('anomaly_count', 0)} line(s) have been flagged for review.",
        size=10, space_after=10)

    kpi_data = [
        ("Total Amount Due",      _kes(s.get("account_total"))),
        ("Pre-Tax Charges",       _kes(s.get("pre_tax_total"))),
        ("Excise Duty (15%)",     _kes(s.get("excise_total"))),
        ("VAT (16%)",             _kes(s.get("vat_total"))),
        ("Outstanding Balance",   _kes(s.get("outstanding_total"))),
        ("Subscriber Lines",      str(s.get("subscriber_count", 0))),
        ("Anomalies Flagged",     str(s.get("anomaly_count", 0))),
        ("Top Division",          data.get("top_division", "—")),
    ]
    widths = [9.5, 6.5]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    _header_row(t, ["Metric", "Value"], widths)
    for i, (metric, value) in enumerate(kpi_data):
        _data_row(t, [metric, value], widths, shade=(i % 2 == 0), bold_first=True)

    # ════════════════════════════════════════════════════════════════════════
    # 2. TAX RECONCILIATION
    # ════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "2.", "Tax Reconciliation")
    pre_tax = float(s.get("pre_tax_total") or 0)
    excise  = float(s.get("excise_total") or 0)
    vat     = float(s.get("vat_total") or 0)
    exp_exc = round(pre_tax * 0.15, 2)
    exp_vat = round((pre_tax + excise) * 0.16, 2)
    exc_diff = round(excise - exp_exc, 2)
    vat_diff = round(vat - exp_vat, 2)

    def diff_str(d):
        if abs(d) < 1: return "✓ Match"
        return f"▲ {_kes(abs(d))} {'over' if d > 0 else 'under'}"

    widths_tax = [5.5, 4.0, 4.0, 3.5]
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    _header_row(t2, ["Component", "Actual (KES)", "Expected (KES)", "Variance"], widths_tax)
    _data_row(t2, ["Pre-Tax Total", f"{pre_tax:,.0f}", "—", "—"],                widths_tax, shade=True,  bold_first=True)
    _data_row(t2, ["Excise Duty (15%)", f"{excise:,.0f}", f"{exp_exc:,.0f}", diff_str(exc_diff)], widths_tax, bold_first=True)
    _data_row(t2, ["VAT (16%)", f"{vat:,.0f}", f"{exp_vat:,.0f}", diff_str(vat_diff)],             widths_tax, shade=True, bold_first=True)
    total = pre_tax + excise + vat
    _data_row(t2, ["Total", f"{total:,.0f}", "—", "—"], widths_tax, bold_first=True)

    # ════════════════════════════════════════════════════════════════════════
    # 3. DIVISION BREAKDOWN
    # ════════════════════════════════════════════════════════════════════════
    if divisions:
        _section_heading(doc, "3.", "Spend by Division / Cost Centre")
        acct_total = float(s.get("account_total") or 1)
        widths_div = [5.5, 2.5, 4.0, 5.0]
        t3 = doc.add_table(rows=1, cols=4)
        t3.style = "Table Grid"
        _header_row(t3, ["Division", "Lines", "Amount (KES)", "Share of Total"], widths_div)
        for i, div in enumerate(divisions):
            pct = f"{(float(div.get('total', 0)) / acct_total * 100):.1f}%"
            bar = "█" * int(float(div.get("total", 0)) / acct_total * 20)
            _data_row(t3, [div.get("division", "—"), str(div.get("count", "")),
                           f"{float(div.get('total', 0)):,.0f}", f"{pct}  {bar}"],
                      widths_div, shade=(i % 2 == 0), bold_first=True)

    # ════════════════════════════════════════════════════════════════════════
    # 4. TOP SUBSCRIBERS
    # ════════════════════════════════════════════════════════════════════════
    if top_subs:
        _section_heading(doc, "4.", "Top Subscribers by Spend")
        widths_sub = [5.5, 3.2, 3.5, 5.3]
        t4 = doc.add_table(rows=1, cols=4)
        t4.style = "Table Grid"
        _header_row(t4, ["Name", "Division", "Amount (KES)", "Tariff Plan"], widths_sub)
        for i, sub in enumerate(top_subs):
            _data_row(t4,
                      [sub.get("raw_name", sub.get("subscriber_number", "—")),
                       sub.get("division", "—"),
                       f"{float(sub.get('amount_due_kes', 0)):,.0f}",
                       sub.get("tariff_plan", "—")],
                      widths_sub, shade=(i % 2 == 0), bold_first=True)

    # ════════════════════════════════════════════════════════════════════════
    # 5. ANOMALIES
    # ════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "5.", "Flagged Anomalies")
    if anomalies:
        widths_anom = [5.0, 3.0, 3.0, 6.5]
        t5 = doc.add_table(rows=1, cols=4)
        t5.style = "Table Grid"
        _header_row(t5, ["Subscriber", "Division", "Amount (KES)", "Reason"], widths_anom)
        for i, anom in enumerate(anomalies):
            _data_row(t5,
                      [anom.get("raw_name", anom.get("subscriber_number", "—")),
                       anom.get("division", "—"),
                       f"{float(anom.get('amount_due_kes', 0)):,.0f}",
                       anom.get("anomaly_reason", "—")],
                      widths_anom, shade=(i % 2 == 0))
    else:
        _add_para(doc, "✓  No anomalies detected in this bill.", size=10,
                  colour=TEAL, space_after=6)

    # ════════════════════════════════════════════════════════════════════════
    # 6. TREND ANALYSIS  (only if multiple bills loaded)
    # ════════════════════════════════════════════════════════════════════════
    if trends and len(trends) >= 2:
        _section_heading(doc, "6.", "Month-on-Month Trend")
        widths_tr = [4.0, 4.0, 3.5, 3.5, 2.5]
        t6 = doc.add_table(rows=1, cols=5)
        t6.style = "Table Grid"
        _header_row(t6, ["Period", "Total (KES)", "vs Prior", "Subscribers", "Anomalies"], widths_tr)
        for i, tr in enumerate(trends):
            prev = trends[i - 1] if i > 0 else None
            if prev:
                delta = float(tr.get("account_total", 0)) - float(prev.get("account_total", 0))
                delta_str = f"{'▲' if delta >= 0 else '▼'} {abs(delta):,.0f}"
            else:
                delta_str = "—"
            _data_row(t6,
                      [tr.get("statement_date", "—"),
                       f"{float(tr.get('account_total', 0)):,.0f}",
                       delta_str,
                       str(tr.get("subscriber_count", "")),
                       str(tr.get("anomaly_count", ""))],
                      widths_tr, shade=(i % 2 == 0))

    # ════════════════════════════════════════════════════════════════════════
    # 7. SUBSCRIBER LIFECYCLE & WASTE
    # ════════════════════════════════════════════════════════════════════════
    waste = data.get("waste") or {}
    if waste.get("dormant_billed") or waste.get("deactivated") or waste.get("top_increases"):
        _section_heading(doc, "7.", "Subscriber Lifecycle & Waste")
        wsum = waste.get("summary", {})
        ref = waste.get("reference_period") or stmt_date
        _add_para(doc,
            f"As of {ref}: {wsum.get('dormant_billed_count', 0)} line(s) billed without any "
            f"usage ({_kes(wsum.get('dormant_billed_kes', 0))} of potentially recoverable "
            f"spend); {wsum.get('deactivated_count', 0)} line(s) dropped off versus "
            f"{waste.get('prev_period') or 'the prior bill'}.",
            size=10, space_after=6)

        dorm = waste.get("dormant_billed", [])[:12]
        if dorm:
            _add_para(doc, "Billed but unused (dormant)", bold=True, size=10,
                      space_before=4, space_after=3)
            wcols = [4.0, 6.5, 4.0, 3.0]
            td = doc.add_table(rows=1, cols=4); td.style = "Table Grid"
            _header_row(td, ["Number", "Name", "Division", "KES"], wcols)
            for i, r in enumerate(dorm):
                _data_row(td, [r.get("display_number", ""), _clip(r.get("raw_name", "")),
                               r.get("division", ""), f"{r.get('amount_kes', 0):,.0f}"],
                          wcols, shade=(i % 2 == 0))

        inc = waste.get("top_increases", [])[:10]
        if inc:
            _add_para(doc, "Largest month-on-month increases", bold=True, size=10,
                      space_before=8, space_after=3)
            wcols = [3.6, 5.4, 2.9, 2.8, 2.8]
            ti = doc.add_table(rows=1, cols=5); ti.style = "Table Grid"
            _header_row(ti, ["Number", "Name", "Prev", "Now", "Change"], wcols)
            for i, r in enumerate(inc):
                _data_row(ti, [r.get("display_number", ""), _clip(r.get("raw_name", "")),
                               f"{r.get('prev_kes', 0):,.0f}", f"{r.get('curr_kes', 0):,.0f}",
                               f"+{r.get('delta_kes', 0):,.0f}"],
                          wcols, shade=(i % 2 == 0), red_last=True)

        deact = waste.get("deactivated", [])[:10]
        if deact:
            _add_para(doc, "Dropped off (possible deactivations)", bold=True, size=10,
                      space_before=8, space_after=3)
            wcols = [4.0, 6.5, 4.0, 3.0]
            tx = doc.add_table(rows=1, cols=4); tx.style = "Table Grid"
            _header_row(tx, ["Number", "Name", "Division", "Last KES"], wcols)
            for i, r in enumerate(deact):
                _data_row(tx, [r.get("display_number", ""), _clip(r.get("raw_name", "")),
                               r.get("division", ""), f"{r.get('last_amount_kes', 0):,.0f}"],
                          wcols, shade=(i % 2 == 0))

    # ════════════════════════════════════════════════════════════════════════
    # 8. RECOMMENDATIONS
    # ════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "8.", "Recommendations")
    recs = []
    if s.get("anomaly_count", 0) > 0:
        recs.append(f"Review {s['anomaly_count']} flagged line(s) — investigate high-spend and unclassified subscribers.")
    if any(d.get("division") == "Unclassified" for d in divisions):
        recs.append("Classify unclassified subscriber lines by adding mapping rules to improve cost centre accuracy.")
    outstanding = float(s.get("outstanding_total") or 0)
    if outstanding > 0:
        recs.append(f"Address outstanding balance of {_kes(outstanding)} to avoid service interruptions.")
    if abs(exc_diff) > 1 or abs(vat_diff) > 1:
        recs.append("Tax reconciliation variance detected — verify excise and VAT calculations with Safaricom.")
    if len(trends) < 3:
        recs.append("Import historical bills to unlock multi-month trend analysis and spend forecasting.")
    recs.append("Review top spenders against approved tariff plans and usage policies.")
    recs.append("Run a quarterly line audit to deactivate unused SIMs and reduce fleet cost.")

    for i, rec in enumerate(recs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i}.  ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE
        run2 = p.add_run(rec)
        run2.font.size = Pt(10)

    # ── Footer note ─────────────────────────────────────────────────────────
    _add_para(doc, space_before=20)
    _add_para(doc,
        f"Report generated by TelecomLens on {generated}. "
        "Data sourced directly from Safaricom postpay invoice PDF. "
        "Confidential — for internal use only.",
        size=8, colour=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Serialise ────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
